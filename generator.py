import re
import docx

"""
doc = docx.Document("learning_agreement.docx")

all_paras = doc.paragraphs

for paragraph in doc.paragraphs: 
    if "{student_first_name}" in paragraph.text: 
        print (paragraph.text) 
        paragraph.text = 'Matej' 

doc.save("new_learning_agreement.docx")

"""

def docx_replace_regex(doc_obj, regex , replace): 
 
    for p in doc_obj.paragraphs: 
        if regex.search(p.text): 
            inline = p.runs 
            # Loop added to work with runs (strings with same style) 
            for i in range(len(inline)): 
                if regex.search(inline[i].text): 
                    text = regex.sub(replace, inline[i].text) 
                    inline[i].text = text
""" 
    for table in doc_obj.tables: 
        for row in table.rows: 
            for cell in row.cells: 
                docx_replace_regex(cell, regex , replace) 
"""
regex1 = re.compile(r"{student_first_name}") 
replace1 = "Matej"

filename = "learning_agreement.docx" 
doc = docx.Document(filename) 
docx_replace_regex(doc, regex1 , replace1) 
doc.save('new_learning_agreement.docx') 