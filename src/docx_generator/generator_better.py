import re
import time

from docx import Document

company_name_regex = re.compile("company_name")
company_address_regex = re.compile("company_address")

cp_first_name_regex = re.compile("cp_first_name")
cp_last_name_regex = re.compile("cp_last_name")
cp_email_regex = re.compile("cp_email")
cp_phone_num_regex = re.compile("cp_phone_num")

student_first_name_regex = re.compile("student_first_name")
student_last_name_regex = re.compile("student_last_name")
student_email_regex = re.compile("student_email")

company = {
    "name": "intercollege",
    "address": "somewhere, 1"
}

contact_person = {
    "first_name": "Milos",
    "last_name": "The goat",
    "email": "milos@thegoat",
    "phone_num": "123456789"
}

student = {
    "first_name": "Laur",
    "last_name": "Bogdan",
    "email": "laur@bogdan"
}


def paragraph_replace_text(paragraph, regex, replace_str):
    """Return `paragraph` after replacing all matches for `regex` with `replace_str`.

    `regex` is a compiled regular expression prepared with `re.compile(pattern)`
    according to the Python library documentation for the `re` module.
    """
    # --- a paragraph may contain more than one match, loop until all are replaced ---
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        # --- when there's a match, we need to modify run.text for each run that
        # --- contains any part of the match-string.
        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        # --- Skip over any leading runs that do not contain the match ---
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        # --- Match starts somewhere in the current run. Replace match-str prefix
        # --- occurring in this run with entire replacement str.
        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len  # --- note this is run-len before replacement ---

        # --- Remove any suffix of match word that occurs in following runs. Note that
        # --- such a suffix will always begin at the first character of the run. Also
        # --- note a suffix can span one or more entire following runs.
        for run in runs:  # --- next and remaining runs, uses same iterator ---
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len

    # --- optionally get rid of any "spanned" runs that are now empty. This
    # --- could potentially delete things like inline pictures, so use your judgement.
    # for run in paragraph.runs:
    #     if run.text == "":
    #         r = run._r
    #         r.getparent().remove(r)

    return paragraph


def generate_learning_goals(doc_path):
    document = Document(doc_path)
    for para in document.paragraphs:
        paragraph_replace_text(para, company_name_regex, company["name"])
        paragraph_replace_text(para, company_address_regex, company["address"])

        paragraph_replace_text(para, cp_first_name_regex, contact_person["first_name"])
        paragraph_replace_text(para, cp_last_name_regex, contact_person["last_name"])
        paragraph_replace_text(para, cp_email_regex, contact_person["email"])
        paragraph_replace_text(para, cp_phone_num_regex, contact_person["phone_num"])

        paragraph_replace_text(para, student_first_name_regex, student["first_name"])
        paragraph_replace_text(para, student_last_name_regex, student["last_name"])
        paragraph_replace_text(para, student_email_regex, student["email"])
    document.save("C:/Users/matko/Desktop/python_docx_generator/generated_docs/learning_agreement_new.docx")
    return document


def async_example():
    print("starting sleep")
    time.sleep(5)
    print("finished sleep")
    return "slept for 5 sec"


"""
if __name__ == "__main__":
    document = Document("learning_agreement.docx")
    for para in document.paragraphs:
        paragraph_replace_text(para, company_name_regex, company["name"])
        paragraph_replace_text(para, company_address_regex, company["address"])

        paragraph_replace_text(para, cp_first_name_regex, contact_person["first_name"])
        paragraph_replace_text(para, cp_last_name_regex, contact_person["last_name"])
        paragraph_replace_text(para, cp_email_regex, contact_person["email"])
        paragraph_replace_text(para, cp_phone_num_regex, contact_person["phone_num"])

        paragraph_replace_text(para, student_first_name_regex, student["first_name"])
        paragraph_replace_text(para, student_last_name_regex, student["last_name"])
        paragraph_replace_text(para, student_email_regex, student["email"])
    document.save("learning_agreement_new.docx")
"""
